"""
Generate Sample DOCX File for Police Duty Card Management System
Creates a realistic .docx file containing section headings and deployment tables with UP Police Hindi duty allocation records.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_sample_docx(output_path="sample_duty_list.docx"):
    doc = Document()

    # Document Header
    heading = doc.add_heading("उत्तर प्रदेश पुलिस ड्यूटी कार्ड / वीआईपी सुरक्षा ड्यूटी आवंटन 2026", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph("जनपद: वाराणसी | कार्यक्रम: श्री काशी विश्वनाथ मंदिर सुरक्षा व्यवस्था")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Section Heading 1 (Zone-01 & Sector-01)
    h_zone1 = doc.add_paragraph()
    run1 = h_zone1.add_run("जोन-01 (मंदिर परिसर एवं मुख्य मार्ग) | सेक्टर-01 (गंगा द्वार बैरियर)")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = RGBColor(11, 25, 44)

    # Table 1 for Zone 1
    table1 = doc.add_table(rows=1, cols=9)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = 'Table Grid'

    headers = ["क्र.सं.", "नाम", "पदनाम", "मोबाईल नंबर", "मूल तैनाती", "जनपद", "ड्यूटी स्थल", "जोन", "सेक्टर"]
    for i, header_text in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header_text
        shading = parse_xml(r'<w:shd {} w:fill="0B192C"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

    records1 = [
        {"sno": "1", "name": "vfer dqekj", "rank": "fu0", "mobile": "9876543210", "posting": "Fkkuk dksrokyh", "district": "Varanasi / वाराणसी", "duty_place": "xaxk }kj - xsV uEcj 1", "zone": "", "sector": ""},
        {"sno": "2", "name": "Rajesh Singh / राजेश सिंह", "rank": "m0fu0", "mobile": "9812345678", "posting": "Police Lines Varanasi", "district": "Varanasi / वाराणसी", "duty_place": "Chowk Crossing / चौक चौराहा", "zone": "", "sector": ""}
    ]

    for rec in records1:
        row_cells = table1.add_row().cells
        row_cells[0].text = rec["sno"]
        row_cells[1].text = rec["name"]
        row_cells[2].text = rec["rank"]
        row_cells[3].text = rec["mobile"]
        row_cells[4].text = rec["posting"]
        row_cells[5].text = rec["district"]
        row_cells[6].text = rec["duty_place"]
        row_cells[7].text = rec["zone"]
        row_cells[8].text = rec["sector"]

    doc.add_paragraph("")

    # Section Heading 2 (Zone-02 & Sector-02)
    h_zone2 = doc.add_paragraph()
    run2 = h_zone2.add_run("जोन-02 (घाट क्षेत्र) | सेक्टर-02 (दशाश्वमेध घाट परिसर)")
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(11, 25, 44)

    # Table 2 for Zone 2
    table2 = doc.add_table(rows=1, cols=9)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'

    for i, header_text in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = header_text
        shading = parse_xml(r'<w:shd {} w:fill="0B192C"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

    records2 = [
        {"sno": "3", "name": "jkes'oj izlkn", "rank": "gs0dk0", "mobile": "9935123456", "posting": "Fkkuk n'kk'oes?k", "district": "Varanasi / वाराणसी", "duty_place": "n'kk'oes?k ?kkV VIP स्टेज", "zone": "", "sector": ""},
        {"sno": "4", "name": "Priyanka Sharma / प्रियंका शर्मा", "rank": "e0dk0", "mobile": "9798765432", "posting": "Women Police Station", "district": "Varanasi / वाराणसी", "duty_place": "काशी विश्वनाथ मंदिर गेट 2 चेकिंग प्वाइंट", "zone": "", "sector": ""},
        {"sno": "5", "name": "lquhy dqekj", "rank": "dk0", "mobile": "9415098765", "posting": "Fkkuk yadjk", "district": "Varanasi / वाराणसी", "duty_place": "BHU Main Gate बैरियर", "zone": "", "sector": ""}
    ]

    for rec in records2:
        row_cells = table2.add_row().cells
        row_cells[0].text = rec["sno"]
        row_cells[1].text = rec["name"]
        row_cells[2].text = rec["rank"]
        row_cells[3].text = rec["mobile"]
        row_cells[4].text = rec["posting"]
        row_cells[5].text = rec["district"]
        row_cells[6].text = rec["duty_place"]
        row_cells[7].text = rec["zone"]
        row_cells[8].text = rec["sector"]

    doc.save(output_path)
    print(f"Sample DOCX duty file created at: {output_path}")

if __name__ == "__main__":
    create_sample_docx()
